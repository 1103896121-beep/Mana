import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional

# 配置日志规范 (符合 Antigravity 规则)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_ai_enhanced_report(output_path: Optional[str] = None) -> str:
    """
    创建 AI 辅助开发成本分析报告。
    
    Args:
        output_path: 报告保存路径。如果为 None，则保存到当前目录下的 'reports' 文件夹。
        
    Returns:
        保存的文件路径。
    """
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('家教平台项目成本分析报告 (AI 辅助开发版)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. AI 辅助研发模式概述
    doc.add_heading('一、 AI 辅助研发模式概述', level=1)
    doc.add_paragraph(
        '在本方案中，我们将深度集成先进 AI（如 Antigravity）作为核心生产力。'
        'AI 不再仅仅是辅助工具，而是承担了从架构设计、代码编写、自动化测试到文档生成的 60%-80% 工作量。'
        '这种模式将传统的“多人协作”转化为“人机协同”，极大地压缩了时间成本和人力财务支出。'
    )
    
    # 2. 传统模式 vs AI 辅助模式成本对比
    doc.add_heading('二、 传统模式 vs AI 辅助模式成本对比', level=1)
    
    table_compare = doc.add_table(rows=1, cols=3)
    table_compare.style = 'Table Grid'
    hdr = table_compare.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '维度', '传统开发模式', 'AI 辅助开发模式'
    
    compare_data = [
        ['起步研发周期', '3 - 6 个月', '1 - 2 个月'],
        ['核心团队规模', '4-6 人 (前后端/UI/测试)', '1-2 人 (高级工程师+AI)'],
        ['初期研发投入', '15w - 30w', '3w - 6w'],
        ['错误/维护成本', '高 (人工排查慢)', '低 (AI 自动扫描与修复)']
    ]
    for d in compare_data:
        row = table_compare.add_row().cells
        for i, v in enumerate(d): row[i].text = v

    # 3. 细化 AI 模式下的综合成本
    doc.add_heading('三、 AI 模式下的综合成本拆解', level=1)
    
    # 3.1 研发成本
    doc.add_heading('3.1 研发与技术工具项', level=2)
    dev_table = doc.add_table(rows=1, cols=3)
    dev_table.style = 'Table Grid'
    hdr = dev_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '项目', '说明', '预算预估 (RMB)'
    
    dev_data = [
        ['AI 协作员/工程师', '具备 AI 使用经验的高级工程师 1 名，负责决策、审核与复杂部署', '每月 2w - 3w'],
        ['AI 工具/API 订阅', '高级 AI API 额度 (Claude/GPT/Antigravity)', '每月 0.1w - 0.2w'],
        ['UI/UX 快速原型', '使用 AI 生成设计稿并由人工微调', '一次性 0.3w'],
        ['服务器基础架设', '由 AI 自动编写部署脚本，减少运维人工', '一次性 0.2w']
    ]
    for d in dev_data:
        row = dev_table.add_row().cells
        for i, v in enumerate(d): row[i].text = v

    # 3.2 基础设施成本 (针对小程序)
    doc.add_heading('3.2 基础设施成本 (云服务)', level=2)
    doc.add_paragraph('利用 AI 进行资源监控与自动扩缩容，优化闲置成本：')
    
    cloud_table = doc.add_table(rows=1, cols=3)
    cloud_table.style = 'Table Grid'
    hdr = cloud_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '组件', 'AI 优化建议', '年度预算 (RMB)'
    
    cloud_data = [
        ['计算资源 (ECS)', '初期 2核4G 精简型，AI 监控负载后手动/自动升配', '0.3w - 0.5w'],
        ['数据库 (Serverless)', '采用按量付费的数据库，AI 负责定时冷备份与清理', '0.2w - 0.4w'],
        ['流量与存储', 'AI 自动压缩图片/视频素材减少 CDN/OSS 消耗', '0.1w']
    ]
    for d in cloud_data:
        row = cloud_table.add_row().cells
        for i, v in enumerate(d): row[i].text = v

    # 3.3 运营端 AI 替代方案
    doc.add_heading('3.3 运营端 AI 替代与降本', level=2)
    doc.add_paragraph('通过集成 AI Agent，可以大幅取代中低级运营岗位：')
    
    ops_table = doc.add_table(rows=1, cols=3)
    ops_table.style = 'Table Grid'
    hdr = ops_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '应用场景', 'AI 替代方案', '节省估算 (RMB)'
    
    ops_data = [
        ['教师资格审核', 'AI OCR 识别证书 + 多源验证 + 逻辑判别', '节省 1 名审核员'],
        ['首推客服', 'AI 智能问答机器人处理 90% 常见咨询', '节省 1 名初级客服'],
        ['推广素材生成', 'AI 批量生成小红书文案与配图', '节省 1 名美工/文案']
    ]
    for d in ops_data:
        row = ops_table.add_row().cells
        for i, v in enumerate(d): row[i].text = v

    # 4. 总结：AI 辅助下的“降本曲线”
    doc.add_heading('四、 总结：AI 辅助下的“降本曲线”', level=1)
    doc.add_paragraph(
        '通过 AI 介入，项目的起步研发成本从 15w+ 降至 5w 左右，降幅超过 60%。'
        '长期运营成本由于 AI 客服和审核的引入，人力薪资支出可减少约 40%-50%。'
        '结论：使用 AI 辅助研发是目前初创项目在资源有限情况下实现“降维打击”的最佳路径。'
    )
    
    # 保存逻辑优化：避免硬编码根路径，优先使用相对路径
    if not output_path:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(base_dir, 'reports')
    else:
        output_dir = os.path.dirname(output_path)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    save_path = output_path or os.path.join(output_dir, '家教平台AI辅助开发成本分析.docx')
    doc.save(save_path)
    logger.info(f"AI enhanced report saved successfully at: {save_path}")
    return save_path

if __name__ == "__main__":
    create_ai_enhanced_report()
