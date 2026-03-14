import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_system_design():
    doc = Document()
    
    # 标题
    title = doc.add_heading('家教平台系统功能设计方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 总体设计思想
    doc.add_heading('一、 总体设计思想', level=1)
    doc.add_paragraph(
        '本系统深度参考“Boss 直聘”的直连模式，取消中间中介环节，通过“老师发布简历”与“家长发布需求”双向驱动。'
        '核心引擎在于【即时通讯】与【精准画像匹配】，确保双方能够像找工作一样通过沟通快速建立信任并达成交易。'
    )
    
    # 2. 核心功能矩阵
    doc.add_heading('二、 核心功能矩阵', level=1)
    
    # 2.1 家长/学生端
    doc.add_heading('2.1 家长/学生端 (需求侧)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('智能找老师：').bold = True
    p.add_run('支持按科目、地理位置(距离)、教龄、老师类别（大学生/专业教师）进行多维筛选。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('发布需求：').bold = True
    p.add_run('家长发布包含补习科目、时间、价格区间、具体地点的需求贴，系统向匹配老师推送通知。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('直连沟通 (IM)：').bold = True
    p.add_run('查看老师 profile 后直接发起对话。支持发送位置、课程表、试课邀请。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('订单管理：').bold = True
    p.add_run('确认授课计费、打卡消课、服务评价。')
    
    # 2.2 老师端
    doc.add_heading('2.2 老师端 (供给侧)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('精美简历：').bold = True
    p.add_run('多媒体简历展示，包含过往教学案例、学历认证、专业证书、教学理念视频。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('抢单中心：').bold = True
    p.add_run('实时查看附近的家长需求，支持一键发送“开场白”和简历。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('日程管理：').bold = True
    p.add_run('内置教务周历，设置忙碌/空闲时间，避免排课冲突。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('钱包中心：').bold = True
    p.add_run('查看应收金额、待解冻课时费、提现申请、纳税说明。')
    
    # 2.3 管理后台
    doc.add_heading('2.3 管理后台 (平台管控)', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('风控审核：').bold = True
    p.add_run('利用 AI OCR 自动核审身份证及学历证，结合人工二次复核，确保师资真实性。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('资金监管：').bold = True
    p.add_run('监控资金池安全性，处理退款争议及平台抽成清算。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('内容审校：').bold = True
    p.add_run('IM 对话敏感词过滤、非法外流联系方式监测（保护平台私域）。')
    
    # 3. 核心业务流程
    doc.add_heading('三、 核心业务流程设计', level=1)
    
    # 流程表
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '流程环节', '关键动作', '数据保障'
    
    flow_data = [
        ['入驻/实名', '老师上传学历证明, 家长实名验证', '公安/学信网接口调用'],
        ['意向建立', '通过筛选/搜索发起聊天', 'IM实时通知系统'],
        ['试课意向', '老师发送试课邀请卡, 家长支付试课订金', '预授权冻结机制'],
        ['正式授课', '家长全额购买课包, 资金平台监管', '课时包模型设计'],
        ['确认结课', '每节课扫描老师打卡二维码, 家长确认', '轨迹匹配与GPS验证'],
        ['佣金分发', '订单完成后, 自动扣除抽成并向老师钱包分发', '子商户分账系统']
    ]
    for d in flow_data:
        row = table.add_row().cells
        for i, v in enumerate(d): row[i].text = v
        
    # 4. 技术方案建议 (AI 辅助视角)
    doc.add_heading('四、 技术架构方案 (AI 辅助实现)', level=1)
    doc.add_paragraph('利用 AI 辅助快速构建的核心组件如下：')
    
    doc.add_paragraph('1. 匹配引擎：利用 AI Embedding 向量化家长需求与老师简历，实现语义级的“兴趣与风格”匹配而非简单的关键词匹配。', style='List Number')
    doc.add_paragraph('2. IM 系统：建议集成第三方成熟 SDK（如环信、网易云信）以缩短开发周期，由 AI 编写协议转换层。', style='List Number')
    doc.add_paragraph('3. 地图 LBS：集成高德/腾讯地图 API，AI 自动处理经纬度计算及路线时间预估。', style='List Number')
    doc.add_paragraph('4. 自动化风控：接入 AI 视觉识别模型，秒级识别非合规证书。', style='List Number')
    
    doc.add_heading('五、 阶段性实施建议', level=1)
    doc.add_paragraph('MVP (最小可行性产品) 阶段：聚焦于“简历发布”与“核心 IM”，暂缓复杂的排课系统。', style='Normal')
    doc.add_paragraph('Beta 阶段：引入“课程包”与“担保支付”闭环。', style='Normal')
    doc.add_paragraph('正式运营：引入 AI 智能推荐算法与全方位的合规审核。', style='Normal')

    output_path = r'E:\workrooten\myself\家教平台系统功能设计方案.docx'
    doc.save(output_path)
    print(f"Design document generated at: {output_path}")

if __name__ == "__main__":
    generate_system_design()
