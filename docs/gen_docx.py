import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Mana (能量) —— 产品需求分析文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 项目背景与愿景
    doc.add_heading('1. 项目背景与愿景', level=1)
    p = doc.add_paragraph()
    p.add_run('“生产力不是清单的堆砌，而是能量的流动。”').bold = True
    doc.add_paragraph('Mana 是一款旨在打破传统“枯燥列表式”任务管理的创新应用。它将抽象的任务管理具象化为“能量管理”，通过物理引擎、魔法化视觉和灵动音效，为用户提供一种掌控感与愉悦感并存的效率体验。')

    # 2. 目标用户
    doc.add_heading('2. 目标用户', level=1)
    doc.add_paragraph('主要面向追求审美、重视心理平衡、活跃于 iOS 生态系统的创造者、职场精英及极简主义者。')

    # 3. 核心功能需求
    doc.add_heading('3. 核心功能需求', level=1)
    
    doc.add_heading('3.1 时间驱动的任务系统', level=2)
    doc.add_paragraph('• 任务时长分配：将原有的能量值概念转化为实用的“所需时间（分钟）”，支持 1-300 分钟设定。')
    doc.add_paragraph('• 可视化进度：界面直观呈现每日积累的产出时间对比，增强掌控感。')
    doc.add_paragraph('• 膨胀与爆炸反馈：任务完成时触发 2-3s 的气泡膨胀动效，模拟压力积聚后释放的物理质感，随后触发粒子爆炸销毁。')
    
    doc.add_heading('3.2 任务列表与排序系统', level=2)
    doc.add_paragraph('• 多维排序：支持按创建时间、所需时长进行升序/降序排列。')
    doc.add_paragraph('• 气泡列表：任务项采用圆润透明的气泡包裹，列表项显示创建的具体时间点。')
    doc.add_paragraph('• 比例优化：针对 iOS 用户调大全局字体与交互组件比例，确保单手操作的准确性。')
    
    doc.add_heading('3.3 动态动效与智能关怀提示', level=2)
    doc.add_paragraph('• 膨胀与爆炸反馈：任务完成时触发 2-3s 的气泡膨胀动效，模拟压力积聚后释放的物理质感。')
    doc.add_paragraph('• 智能提醒触发方案：系统根据以下三个维度实时监测并随机弹出关怀提示词：')
    doc.add_paragraph('  1. 累计数量：每日累计每完成 5 个任务；')
    doc.add_paragraph('  2. 累计时长：每日累计任务总时长达到 120 分钟或其倍数；')
    doc.add_paragraph('  3. 高频强度：30 分钟内连续完成 3 项及以上任务。')
    
    doc.add_heading('3.4 艺术化产出分享', level=2)
    doc.add_paragraph('• 定制卡片：将一天的生产力数据转化为极致简约的视觉卡片，便于社交分享。')

    # 4. 视觉语言与多主题
    doc.add_heading('4. 视觉语言与多主题支持', level=1)
    doc.add_paragraph('• 关键词：灵动、空灵、平衡、双相。')
    doc.add_paragraph('• 双色模式：支持亮色 (Light) 与暗色 (Dark) 两套视觉系统，满足全天候使用场景。')
    doc.add_paragraph('• 设计元素：玻璃态 (Glassmorphism)、微动效 (Micro-animations)、大气泡感。')

    # 5. 技术架构与交互要求
    doc.add_heading('5. 技术架构与交互要求', level=1)
    doc.add_paragraph('• 架构方案：无后端纯前端架构 (Serverless Frontend-Only)。系统完全运行在用户苹果手机本地，无需外部服务器支持。')
    doc.add_paragraph('• 数据持久化：利用浏览器 LocalStorage 或 IndexedDB 实现本地持久化方案，确保用户数据安全存储于设备本地且支持离线使用。')
    doc.add_paragraph('• 响应式设计：适配 iOS 各种屏幕尺寸及灵动岛等硬件特性。')
    doc.add_paragraph('• 动画引擎：底层需结合 Framer Motion 或 Canvas 实现复杂的物理粒子效果。')
    doc.add_paragraph('• 交互性能：保持 60/120 FPS 的极致流畅度，满足用户感官预期。')

    # Save
    output_path = os.path.join('e:\\workrooten\\Mana\\docs', 'Mana-产品需求分析文档.docx')
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_doc()
