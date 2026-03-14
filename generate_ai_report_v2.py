import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ai_report_v2():
    doc = Document()
    
    # 标题
    title = doc.add_heading('家教平台项目可行性研究报告 (AI 辅助研发版)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 品牌命名建议
    doc.add_heading('一、 品牌命名建议 (特色化与独特性)', level=1)
    doc.add_paragraph('根据“Boss直聘式”直连逻辑与教育行业属性，建议以下品牌名：')
    
    name_table = doc.add_table(rows=1, cols=3)
    name_table.style = 'Table Grid'
    hdr = name_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '建议名称', '品牌寓意', '特色分'
    
    names = [
        ['见师 (JianShi)', '“见”即“面谈/直聘”，“师”即良师。寓意家长能直接见到真实、优质的老师，去中介化。', '★★★★★'],
        ['师友 (ShiYou)', '既是老师，也是朋友。强调“Boss”般的平权沟通感，打破传统补习的压迫感。', '★★★★☆'],
        ['启明星 (StarGuide)', '寓意在学习的道路上，名师如晨星指引方向，具有极强的品牌辨识度和正面联想。', '★★★★★'],
        ['学搭 (XueDa)', '“学习搭子”。符合当下大群体及社交趋势，定位精准、亲和力强，易于在社群传播。', '★★★★☆'],
        ['良师约 (GoodTeacherLink)', '简单直白，突出“约”的动作，强调预约制与直约感，信任度极高。', '★★★★']
    ]
    for n in names:
        row = name_table.add_row().cells
        for i, v in enumerate(n): row[i].text = v

    # 2. AI 辅助研发模式及成本重核
    doc.add_heading('二、 AI 辅助研发模式下的降本增效分析', level=1)
    doc.add_paragraph(
        '引入 Antigravity 等先进 AI 辅助后，研发逻辑从“堆人力”转向“高杠杆”。'
        'AI 可承担 80% 的代码编写、自动化测试及 API 文档生成。'
    )
    
    # 2.1 研发成本对比
    doc.add_heading('2.1 综合研发成本对比', level=2)
    compare_table = doc.add_table(rows=1, cols=3)
    compare_table.style = 'Table Grid'
    hdr = compare_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '项目', '传统开发模式', 'AI 辅助开发模式'
    
    costs = [
        ['起步研发周期', '120 - 180 天', '30 - 45 天'],
        ['核心人力成本', '20.0w - 40.0w (5人团队)', '3.0w - 6.0w (1人+AI)'],
        ['API 与工具支出', '~ 0.5w', '0.5w - 1.2w (含高性能AI额度)'],
        ['总研发投入', '20.5w - 40.5w', '3.5w - 7.2w']
    ]
    for c in costs:
        row = compare_table.add_row().cells
        for i, v in enumerate(c): row[i].text = v

    # 2.2 AI 在运营端的减负
    doc.add_heading('2.2 AI 在运营环节的长期降本', level=2)
    ops_table = doc.add_table(rows=1, cols=3)
    ops_table.style = 'Table Grid'
    hdr = ops_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '运营版块', 'AI 应用方式', '人力节省预估'
    
    ops = [
        ['教师资质审核', 'AI OCR 识别证书 + 多级逻辑判定', '节省 1.5 名审核员'],
        ['在线客服', 'AI Agent 实时处理 85% 常见问题', '节省 2 名初级客服'],
        ['营销内容产出', 'AI 生成多套小红书/抖音文案与配图', '节省 1 名美工/文案']
    ]
    for o in ops:
        row = ops_table.add_row().cells
        for i, v in enumerate(o): row[i].text = v

    # 3. 基础设施成本 (AI 优化版)
    doc.add_heading('三、 基础设施成本 (AI 自动监控版)', level=1)
    doc.add_paragraph('利用 AI 监控系统负载并自动调整配置，避免资源浪费。')
    
    infra_table = doc.add_table(rows=1, cols=3)
    infra_table.style = 'Table Grid'
    hdr = infra_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '组件', 'AI 建议规格', '年度预估 (RMB)'
    
    infra = [
        ['应用网关与计算', 'ECS 突发型实例 (AI 动态升配)', '0.4w - 0.6w'],
        ['数据库存储', 'Serverless 模式按量计费 (AI 自动备份)', '0.2w - 0.3w'],
        ['带宽与CDN', '按流量计费 (AI 自动压缩静态资源)', '0.1w - 0.2w']
    ]
    for f in infra:
        row = infra_table.add_row().cells
        for i, v in enumerate(f): row[i].text = v

    # 4. 结论：AI 时代的创业路径
    doc.add_heading('四、 结论：AI 时代的创业路径', level=1)
    doc.add_paragraph(
        '通过深度结合 AI 开发与运营，项目初期启动资金可从 40w+ 压缩至 10w 以内。'
        '重点投入应转向【营销推广】与【师资品质控】，而非昂贵的【通用技术研发】。'
    )

    output_path = r'E:\workrooten\myself\家教平台AI辅助开发成本分析版.docx'
    doc.save(output_path)
    print(f"AI enhanced report saved at: {output_path}")

if __name__ == "__main__":
    generate_ai_report_v2()
