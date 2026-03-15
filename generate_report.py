import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional

# 配置日志规范
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_report(output_path: Optional[str] = None) -> str:
    """
    创建家教平台项目可行性研究报告。
    
    Args:
        output_path: 报告保存路径。
        
    Returns:
        保存的文件路径。
    """
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('家教平台项目可行性研究报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 项目概述
    doc.add_heading('一、 项目概述', level=1)
    doc.add_paragraph(
        '本项目旨在打造一个高效、透明的家教资源撮合平台（类似“家教版 Boss 直聘”）。'
        '平台连接两端核心用户：需要个性化辅导的家庭用户（家长）和拥有专业教学能力的家教老师（大学生、在职教师、机构老师）。'
        '通过数字化精准匹配、担保交易及评价体系，解决传统家教市场信息不对称、信任成本高、匹配效率低等痛点。'
    )
    
    # 2. 市场分析与可行性
    doc.add_heading('二、 市场分析与可行性', level=1)
    doc.add_paragraph('在当前国内政策与社会环境下，家教平台具有显著的生存空间与增长潜力：')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('政策合规性：').bold = True
    p.add_run('在“双减”政策背景下，学科类培训受到限制，但非学科类（素质教育、艺术、体育等）需求依然旺盛。平台可通过引导教师提供多元化素质服务，确保经营合规。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('供需平衡：').bold = True
    p.add_run('大学生群体有强烈的兼职勤工俭学需求，而家长端对高质量、个性化、上门式辅导的需求始终存在，尤其是针对学生兴趣培养与习惯养成。')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('技术驱动：').bold = True
    p.add_run('利用类似 Boss 直聘的即时沟通模式，让家长与老师直接对话，极大缩短决策链路。')
    
    # 3. 应用形式分析建议
    doc.add_heading('三、 应用形式分析与建议', level=1)
    doc.add_paragraph('针对网站、App、小程序三种形式，对比分析如下：')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '小程序'
    hdr_cells[2].text = '网站 (Web)'
    hdr_cells[3].text = '移动端 App'
    
    data = [
        ['开发成本', '低', '中', '高'],
        ['获客难度', '低 (微信生态转发)', '中 (依赖搜索)', '高 (需下载安装)'],
        ['用户留存', '中 (即用即走)', '低', '高 (常驻手机)'],
        ['功能表现', '中 (受限于平台)', '中', '极佳 (深度性能)'],
        ['适用场景', '初期验证与快速获客', '信息展示与SEO', '深度学习与重度交互']
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph('\n')
    doc.add_paragraph('【建议方案】：', style='Normal').runs[0].bold = True
    doc.add_paragraph(
        '现阶段建议优先以“微信小程序”作为核心入口。理由：家庭用户和大学生均为微信重度用户，小程序无需下载、分享便捷、'
        '闭环支付成熟，是目前获客成本最低、验证逻辑最快的选择。后期流量稳定后再考虑推出 App 以增强用户粘性。'
    )
    
    # 4. 商业模式与抽成
    doc.add_heading('四、 商业模式与收益方案', level=1)
    doc.add_paragraph('平台采用“信息撮合 + 交易担保”模式，通过以下方式获利：')
    doc.add_paragraph('1. 课时费抽成：平台作为第三方担保支付方，抽取每笔交易额的 5%-15%（根据项目类型及老师级别浮动）。', style='List Number')
    doc.add_paragraph('2. 会员/增值服务：老师端可购买“简历加亮”、“优先推荐”等类似 Boss 直聘的功能。', style='List Number')
    doc.add_paragraph('3. 认证费/保证金：对教师进行入驻实名及学历认证，收取的少量行政成本或信誉保证金。', style='List Number')
    
    # 5. 成本分析
    doc.add_heading('五、 成本分析 (初期预算)', level=1)
    doc.add_paragraph('预计初期投入分为以下几个版块：')
    
    cost_data = [
        ['类目', '预估内容', '预估金额 (RMB)'],
        ['技术研发', '小程序开发、服务器架设、支付集成', '5w - 15w (视外包或自建)'],
        ['运营成本', '师资审核、客服支持、日常维护', '每月 1w - 3w'],
        ['营销推广', '校园大使、社交媒体广告、地推物料', '初期 5w - 10w'],
        ['资质合规', '公司注册、ICP备案、法律顾问', '0.5w - 2w']
    ]
    
    table_cost = doc.add_table(rows=1, cols=3)
    table_cost.style = 'Table Grid'
    for i, text in enumerate(cost_data[0]):
        table_cost.rows[0].cells[i].text = text
        
    for row_data in cost_data[1:]:
        row_cells = table_cost.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            
    # 6. 小程序专项成本方案（深度分析）
    doc.add_heading('六、 小程序专项成本方案 (深度分析)', level=1)
    doc.add_paragraph('鉴于建议优先采用小程序形式，现对其成本构成进行精细化拆解：')
    
    # 6.1 研发成本
    doc.add_heading('6.1 研发成本 (一次性)', level=2)
    dev_table = doc.add_table(rows=1, cols=3)
    dev_table.style = 'Table Grid'
    hdr = dev_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '模块/角色', '工作内容', '预估成本 (RMB)'
    
    dev_data = [
        ['前端开发 (Uni-app)', '双端(家长/教师)小程序界面、交互逻辑、支付唤起', '2w - 4w'],
        ['后端开发 (FastAPI)', 'API接口、匹配引擎、数据库设计、安全校验', '3w - 5w'],
        ['UI/UX 设计', '原型图、高保真视觉设计、Logo与品牌规范', '0.8w - 1.5w'],
        ['QA 测试', '功能测试、并发压力测试、多机型兼容性测试', '0.5w - 1w']
    ]
    for d in dev_data:
        row = dev_table.add_row().cells
        for i, v in enumerate(d): row[i].text = v
    doc.add_paragraph('注：以上为初创标准（外包或兼职团队），若自有全职团队成本约在 15w-25w 之间。')
 
    # 保存逻辑
    if not output_path:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(base_dir, 'reports')
    else:
        output_dir = os.path.dirname(output_path)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    save_path = output_path or os.path.join(output_dir, '家教平台可行性研究报告.docx')
    doc.save(save_path)
    logger.info(f"Report updated successfully at: {save_path}")
    return save_path

if __name__ == "__main__":
    create_report()
